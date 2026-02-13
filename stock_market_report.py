import os
import time
import requests
import pandas as pd
import yfinance as yf
import mplfinance as mpf
from fpdf import FPDF
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
import io
import matplotlib.pyplot as plt
import subprocess
import glob
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import shutil

# Import the new modules
from bulletin import create_filtered_market_bulletin
from stocksnews import create_stocks_bulletin
from currency import get_currency_exchange_rates
from fii_dii_data import generate_fii_dii_summary, get_fii_dii_chart
from gainerslosers import get_nifty50_movers
from generate_graph import generate_nifty_report_with_slider
from generate_report import generate_nifty_report_pdf
from get_nifty_pcr import get_pcr_data, get_pcr_trend_chart
from global import get_global_indices_data
from gold import get_chennai_gold_rates
from silver import get_chennai_silver_rates
from heatmap import get_tradingview_heatmap_price
from nifty_oi import get_nifty_oi_data_and_chart
from sgx import get_sgx_nifty_snapshot
from vix import get_vix_data_and_chart

# --- Global Constants for File Names ---
PDF_REPORT_FILENAME = "YouGrow_Daily_Market_Report.pdf"
DOCX_BULLETIN_FILENAME = "Market_Bulletin.docx"

# --- PDF Generation Class ---
class PDF(FPDF):
    def header(self):
        # Logo (if exists)
        # self.image('logo.png', 10, 8, 33)
        self.set_font('Arial', 'B', 15)
        # Move to the right
        self.cell(80)
        # Title
        self.cell(30, 10, 'YouGrow Daily Market Report', 0, 0, 'C')
        # Line break
        self.ln(20)

    def footer(self):
        # Position at 1.5 cm from bottom
        self.set_y(-15)
        # Arial italic 8
        self.set_font('Arial', 'I', 8)
        # Page number
        self.cell(0, 10, 'Page ' + str(self.page_no()) + '/{nb}', 0, 0, 'C')

    def chapter_title(self, num, label):
        self.set_font('Arial', 'B', 12)
        self.set_fill_color(200, 220, 255)
        self.cell(0, 6, 'Chapter %d : %s' % (num, label), 0, 1, 'L', 1)
        self.ln(4)

    def chapter_body(self, body):
        self.set_font('Times', '', 12)
        self.multi_cell(0, 5, body)
        self.ln()

    def add_section_title(self, title):
        self.set_font('Arial', 'B', 14)
        self.set_fill_color(230, 230, 230)
        self.cell(0, 8, title, 0, 1, 'L', 1)
        self.ln(4)
        
    def add_chart_image(self, image_path, w=180, h=100, x=None):
        if os.path.exists(image_path):
            if x:
                self.image(image_path, x=x, w=w, h=h)
            else:
                self.image(image_path, w=w, h=h)
            self.ln(5)
        else:
            self.set_font('Arial', 'I', 10)
            self.cell(0, 10, f"Chart image not found: {image_path}", 0, 1)

# --- Helper Function for PDF Table ---
def add_table_to_pdf(pdf, headers, data, col_widths):
    pdf.set_font('Arial', 'B', 10)
    for header, width in zip(headers, col_widths):
        pdf.cell(width, 7, header, 1, 0, 'C')
    pdf.ln()
    
    pdf.set_font('Arial', '', 9)
    for row in data:
        for item, width in zip(row, col_widths):
            # Ensure item is string
            text = str(item)
            # Handle potentially long text by truncating or multiline (simplified here)
            pdf.cell(width, 7, text, 1, 0, 'C')
        pdf.ln()

# --- Main Report Orchestration Function ---
def generate_full_report():
    print("\n" + "="*50)
    print("STARTING FULL REPORT GENERATION")
    print("="*50)
    
    # Initialize data dictionary to store results
    report_data = {}
    
    # Create a temporary directory for chart images
    if not os.path.exists("temp_charts"):
        os.makedirs("temp_charts")
    
    # --- 1. NIFTY 50 Analysis & Chart ---
    print("\n[1/12] Generating Nifty 50 Report...")
    # This generates "Nifty_Report_Slider.pdf"
    nifty_success = generate_nifty_report_with_slider("temp_charts/nifty_slider.pdf")
    if nifty_success:
        report_data['nifty_slider_pdf'] = "temp_charts/nifty_slider.pdf"
        print("  > Nifty Slider Report generated.")
    else:
        print("  > Failed to generate Nifty Slider Report.")

    # --- 2. NIFTY PCR Analysis ---
    print("\n[2/12] Fetching Nifty PCR Data...")
    pcr_val = get_pcr_data()
    pcr_chart_success = get_pcr_trend_chart("temp_charts/pcr_chart.png")
    if pcr_val:
        report_data['pcr'] = pcr_val
        print(f"  > PCR Value: {pcr_val}")
    if pcr_chart_success:
        report_data['pcr_chart'] = "temp_charts/pcr_chart.png"

    # --- 3. India VIX ---
    print("\n[3/12] Fetching India VIX Data...")
    vix_val, vix_chart_success = get_vix_data_and_chart("temp_charts/vix_chart.png")
    if vix_val:
        report_data['vix'] = vix_val # (current, change, pct_change)
        print(f"  > VIX Value: {vix_val[0]}")
    if vix_chart_success:
        report_data['vix_chart'] = "temp_charts/vix_chart.png"

    # --- 4. FII/DII Data ---
    print("\n[4/12] Fetching FII/DII Data...")
    fii_dii_summary = generate_fii_dii_summary()
    fii_dii_chart_success = get_fii_dii_chart("temp_charts/fii_dii_chart.png")
    if fii_dii_summary:
        report_data['fii_dii'] = fii_dii_summary
        print("  > FII/DII Summary fetched.")
    if fii_dii_chart_success:
        report_data['fii_dii_chart'] = "temp_charts/fii_dii_chart.png"

    # --- 5. SGX Nifty Snapshot ---
    print("\n[5/12] Capturing SGX Nifty Snapshot...")
    sgx_success = get_sgx_nifty_snapshot("temp_charts/sgx_nifty.png")
    if sgx_success:
        report_data['sgx_nifty_chart'] = "temp_charts/sgx_nifty.png"
        print("  > SGX Nifty Snapshot captured.")

    # --- 6. Global Indices ---
    print("\n[6/12] Fetching Global Indices...")
    global_indices = get_global_indices_data()
    if global_indices:
        report_data['global_indices'] = global_indices
        print("  > Global Indices data fetched.")

    # --- 7. Nifty OI Analysis ---
    print("\n[7/12] Fetching Nifty OI Data...")
    nifty_oi = get_nifty_oi_data_and_chart("temp_charts/nifty_oi_chart.png")
    if nifty_oi:
        report_data['nifty_oi'] = nifty_oi
        print("  > Nifty OI data fetched.")

    # --- 8. Gold Rates ---
    print("\n[8/12] Fetching Gold Rates...")
    gold_data = get_chennai_gold_rates()
    if gold_data:
        report_data['gold'] = gold_data
        print("  > Gold Rates fetched.")

    # --- 9. Silver Rates ---
    print("\n[9/12] Fetching Silver Rates...")
    silver_data = get_chennai_silver_rates()
    if silver_data:
        report_data['silver'] = silver_data
        print("  > Silver Rates fetched.")

    # --- 10. Currency Rates ---
    print("\n[10/12] Fetching Currency Rates...")
    currency_data = get_currency_exchange_rates()
    if currency_data:
        report_data['currency'] = currency_data
        print("  > Currency Rates fetched.")

    # --- 11. Market News Bulletin (DOCX) ---
    print("\n[11/12] Generating Market News Bulletin...")
    news_text = create_filtered_market_bulletin(DOCX_BULLETIN_FILENAME)
    if news_text:
        report_data['news_bulletin'] = news_text
        print(f"  > Market Bulletin DOCX saved as {DOCX_BULLETIN_FILENAME}")
    
    # --- 11b. Stocks News Bulletin (DOCX) ---
    print("\n[11b/12] Generating Stocks News Bulletin...")
    stocks_news_success = create_stocks_bulletin("Key_Stocks_to_Watch.docx")
    if stocks_news_success:
        print(f"  > Stocks News Bulletin DOCX saved as Key_Stocks_to_Watch.docx")

    # --- 12. Nifty 50 Heatmap ---
    print("\n[12/12] Capturing Nifty 50 Heatmap...")
    heatmap_success = get_tradingview_heatmap_price("temp_charts/nifty_heatmap.png")
    if heatmap_success:
        report_data['heatmap'] = "temp_charts/nifty_heatmap.png"
        print("  > Heatmap captured.")

    # --- COMPILE MASTER PDF ---
    print("\n" + "="*50)
    print("COMPILING MASTER PDF REPORT")
    print("="*50)

    pdf = PDF()
    pdf.alias_nb_pages()
    
    # Page 1: Executive Summary & Nifty Pulse
    pdf.add_page()
    pdf.add_section_title("1. Market Pulse: Nifty 50")
    
    # We can assume Nifty Slider PDF has been generated. 
    # Ideally, we would merge the PDFs, but for this simple version, 
    # we might capture it as an image or just refer to the data.
    # To keep it simple, we will reuse the logic from generate_graph.py 
    # to find supports and draw them here, or better yet, assume the user
    # wants the single PDF merged. 
    # For now, let's just insert the Nifty Slider Chart if we can specificallly
    # save just that chart image in generate_graph or just reiterate the key points.
    
    # Since generate_nifty_report_with_slider creates a full page PDF, 
    # we will append it at the end or use pypdf to merge.
    # Let's stick to adding content to our FPDF object for now.
    
    if 'nifty_oi' in report_data:
        data = report_data['nifty_oi']
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, f"Spot Price: {data['spot_price']:,.2f}", 0, 1)
        pdf.cell(0, 10, f"Total Call OI: {data['total_calls_oi']}", 0, 1)
        pdf.cell(0, 10, f"Total Put OI: {data['total_puts_oi']}", 0, 1)
        
        # Calculate PCR from OI data if available
        try:
            call_oi_num = float(data['total_calls_oi'].replace('Cr', '').replace('L', '').strip())
            put_oi_num = float(data['total_puts_oi'].replace('Cr', '').replace('L', '').strip())
            # Basic approximation, real calculation is more complex
        except:
            pass

    if 'pcr' in report_data:
        pdf.ln(5)
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, f"Nifty PCR: {report_data['pcr']}", 0, 1)
        
        if 'pcr_chart' in report_data:
             pdf.add_chart_image(report_data['pcr_chart'], w=160, h=90)

    if 'vix' in report_data:
        curr, chg, pct = report_data['vix']
        pdf.ln(5)
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, f"India VIX: {curr} ({chg:+.2f}, {pct:+.2f}%)", 0, 1)
        
        if 'vix_chart' in report_data:
            pdf.add_chart_image(report_data['vix_chart'], w=160, h=90)

    # Page 2: FII/DII & SGX
    pdf.add_page()
    pdf.add_section_title("2. Institutional Activity (FII/DII)")
    
    if 'fii_dii' in report_data:
        # Create table data
        headers = ['Date', 'FII Net', 'DII Net']
        table_data = []
        for row in report_data['fii_dii']:
            table_data.append([row['Date'], row['FII'], row['DII']])
        
        add_table_to_pdf(pdf, headers, table_data, [40, 50, 50])
        pdf.ln(10)
        
    if 'fii_dii_chart' in report_data:
        pdf.add_chart_image(report_data['fii_dii_chart'], w=180, h=100)

    # SGX Nifty
    pdf.ln(5)
    pdf.add_section_title("3. SGX Nifty Snapshot")
    if 'sgx_nifty_chart' in report_data:
        pdf.add_chart_image(report_data['sgx_nifty_chart'], w=180, h=120)

    # Page 3: Global Indices & Currency
    pdf.add_page()
    pdf.add_section_title("4. Global Market Indices")
    
    if 'global_indices' in report_data:
        headers = ['Index', 'LTP', 'Change', '% Change']
        table_data = []
        for idx in report_data['global_indices']:
            table_data.append([
                idx['Name'], 
                f"{idx['LTP']:,.2f}", 
                f"{idx['Change']:+.2f}", 
                f"{idx['Change %']:+.2f}%"
            ])
        add_table_to_pdf(pdf, headers, table_data, [50, 40, 40, 40])
        pdf.ln(10)

    pdf.add_section_title("5. Currency Exchange Rates (vs INR)")
    if 'currency' in report_data:
        headers = ['Currency', 'Name', 'Value (INR)']
        table_data = []
        for curr in report_data['currency']:
            table_data.append([curr['Code'], curr['Name'], f"{curr['Value']:.2f}"])
        add_table_to_pdf(pdf, headers, table_data, [30, 80, 40])
    
    # Page 4: Commodities (Gold/Silver)
    pdf.add_page()
    pdf.add_section_title("6. Commodities: Gold & Silver (Chennai)")
    
    if 'gold' in report_data:
        gold = report_data['gold']
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(0, 10, "Today's Gold Rate (22K): " + f"{gold['today_22k']['price']:,.2f}", 0, 1)
        
        pdf.set_font('Arial', 'I', 10)
        pdf.cell(0, 10, "Historical Gold Rates (Last 10 Days)", 0, 1)
        
        headers = ['Date', '24K Price', '22K Price']
        table_data = []
        for row in gold['last_10_days']:
            table_data.append([row['date'], row['price_24k'], row['price_22k']])
        add_table_to_pdf(pdf, headers, table_data, [40, 50, 50])
        pdf.ln(10)

    if 'silver' in report_data:
        silver = report_data['silver']
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(0, 10, "Today's Silver Rate (1Kg): " + f"{silver['today_per_kg']['price']:,.2f}", 0, 1)
        
        pdf.set_font('Arial', 'I', 10)
        pdf.cell(0, 10, "Historical Silver Rates (Last 10 Days)", 0, 1)
        
        headers = ['Date', '10g Price', '1Kg Price']
        table_data = []
        for row in silver['last_10_days']:
            table_data.append([row['date'], row['price_10g'], row['price_1kg']])
        add_table_to_pdf(pdf, headers, table_data, [40, 50, 50])

    # Page 5: Heatmap
    if 'heatmap' in report_data:
        pdf.add_page()
        pdf.add_section_title("7. Nifty 50 Heatmap")
        # Rotate page to landscape effectively by just placing image large
        # Or just keep portrait and fit image
        pdf.add_chart_image(report_data['heatmap'], w=190, h=140)

    # Page 6: News Headlines
    pdf.add_page()
    pdf.add_section_title("8. Top Market News Headlines")
    
    if 'news_bulletin' in report_data:
        # The news_bulletin is a string of text from the DOCX function
        # We need to split it carefully
        # Simple implementation:
        # Note: The DOCX function returns a string now
        
        news_lines = report_data['news_bulletin'].split('\n')
        pdf.set_font('Times', '', 11)
        for line in news_lines:
            if line.strip():
                # Handling basic encoding issues
                clean_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 6, clean_line)
                pdf.ln(2)

    # --- Save Master PDF ---
    try:
        pdf.output(PDF_REPORT_FILENAME, 'F')
        print(f"\nSUCCESS: Master Report generated as '{PDF_REPORT_FILENAME}'")
    except Exception as e:
        print(f"Error saving PDF: {e}")

    # --- Merge with Nifty Slider PDF (Optional/Advanced) ---
    # Attempt to merge if pyPDF2 or similar is available. 
    # For now, we will leave them as separate files or rely on the image embedding above 
    # (if we had saved the slider as an image).
    
    # Cleanup temp charts
    # shutil.rmtree("temp_charts") 
    print("Temporary charts kept in 'temp_charts/' for debugging.")

if __name__ == "__main__":
    generate_full_report()
